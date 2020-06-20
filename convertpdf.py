import fpdf
from docx import Document

with open('rama.txt') as f1:
    doc= f1.read()
#doc = Document(st)
pdf=fpdf.FPDF()
pdf.add_page()
pdf.set_font("Arial","B",12)
pdf.set_xy(10,0)
#pdf.multi_cell(175,10,doc,border=0,align='J')
pdf.cell(0,20,doc)
pdf.output("newconverted21.pdf")
'''
doc = Document('google.docx')
all_paras=doc.paragraphs
pdf=fpdf.FPDF()
pdf.add_page()
pdf.set_font("Arial","B",12)
pdf.set_xy(10,0)
for para in all_paras:
    pdf.multi_cell(175,10,para.text,border=0,align='J')
#pdf.cell(0,0,con)
pdf.output("newconverted11.pdf")
'''
